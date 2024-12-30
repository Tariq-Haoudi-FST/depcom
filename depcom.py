import streamlit as st
from streamlit_option_menu import option_menu
import io
from matplotlib import pyplot as plt
import streamlit as st
import plotly.express as px
import pandas as pd
import tempfile
import matplotlib.pyplot as plt
import streamlit as st
from st_aggrid import AgGrid, GridOptionsBuilder
from docx import Document
from docx.shared import Inches

from docx import Document
from docx import Document

def create_sales_team_report(df, sales_teams):
    # Créer un document Word
    doc = Document()
    
    # Ajouter un titre
    doc.add_heading('Rapport des Données Groupées par Équipes Commerciales', level=1)

    for team in sales_teams:
        doc.add_heading(f'Données pour l\'Équipe : {team}', level=2)
        
        # Filtrer les données pour l'équipe actuelle
        filtered_df = df[df['SALES TEAM'] == team]
        
        # Regrouper les données
        grouped_df = (
            filtered_df
            .groupby(['OPPORTUNITY TYPE', 'PREVIOUS PIPELINE STAGE'])
            .size()
            .reset_index(name='COUNT')
        )
        
        # Ajouter une colonne 'SUM' pour la somme par OPPORTUNITY TYPE
        grouped_df['SUM'] = grouped_df.groupby('OPPORTUNITY TYPE')['COUNT'].transform('sum')
        
        # Ajouter une ligne pour le total général
        total_row = pd.DataFrame({
            'OPPORTUNITY TYPE': ['TOTAL'],
            'PREVIOUS PIPELINE STAGE': [''],
            'COUNT': [grouped_df['COUNT'].sum()],
            'SUM': [grouped_df['SUM'].sum()]
        })
        grouped_df = pd.concat([grouped_df, total_row], ignore_index=True)

        # Ajouter un tableau pour les données groupées de l'équipe
        table = doc.add_table(rows=1, cols=len(grouped_df.columns))
        table.style = 'Table Grid'
        
        # Ajouter les en-têtes de colonnes
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(grouped_df.columns):
            hdr_cells[i].text = str(column)
        
        # Ajouter les données du DataFrame groupé
        for index, row in grouped_df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    # Enregistrer le document
    doc.save('rapport_pipeline_teams.docx')

from docx import Document
from docx.shared import Inches

from docx import Document
from docx.shared import Inches

def create_word_report2(grouped_df, dfC, unique_villes):
    # Créer un document Word
    doc = Document()
    
    # Ajouter un titre
    doc.add_heading('Rapport des Données Groupées et Filtrées', level=1)

    # Ajouter le tableau groupé
    doc.add_heading('Données Groupées', level=2)
    table_grouped = doc.add_table(rows=1, cols=len(grouped_df.columns))
    table_grouped.style = 'Table Grid'
    
    # Ajouter les en-têtes de colonnes
    hdr_cells = table_grouped.rows[0].cells
    for i, column in enumerate(grouped_df.columns):
        hdr_cells[i].text = str(column)
    
    # Ajouter les données du DataFrame groupé
    for index, row in grouped_df.iterrows():
        row_cells = table_grouped.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)

    # Ajouter une section pour les données filtrées par ville
    for ville in unique_villes:
        doc.add_heading(f'Données Filtrées pour {ville}', level=2)
        
        # Filtrer les données pour la ville actuelle
        filtered_df = dfC[dfC['VILLE'] == ville]
        
        # Ajouter un tableau pour la ville
        table_ville = doc.add_table(rows=1, cols=len(filtered_df.columns))
        table_ville.style = 'Table Grid'
        
        # Ajouter les en-têtes de colonnes pour la ville
        hdr_cells_ville = table_ville.rows[0].cells
        for i, column in enumerate(filtered_df.columns):
            hdr_cells_ville[i].text = str(column)
        
        # Ajouter les données du DataFrame de la ville
        for index, row in filtered_df.iterrows():
            row_cells = table_ville.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

        # Calculer et ajouter le total du CA pour la ville
        total_ca_ville = filtered_df['CA'].sum()
        doc.add_paragraph(f"Total du CA pour {ville} : {total_ca_ville}")

    # Enregistrer le document
    doc.save('rapport_donnees1.docx')
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import io
import plotly.express as px

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import io
import plotly.express as px

from docx import Document
from docx.shared import Inches

def create_word_report1(grouped_df):
    # Créer un document Word
    doc = Document()
    
    # Ajouter un titre
    doc.add_heading('Rapport des Données Groupées', level=1)

    # Ajouter une description
    doc.add_paragraph('Ce rapport présente les données groupées par VILLE, COMMERCIAL (E) et ETAT.')

    # Ajouter un tableau
    table = doc.add_table(rows=1, cols=len(grouped_df.columns))
    table.style = 'Table Grid'
    
    # Ajouter les en-têtes de colonnes
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(grouped_df.columns):
        hdr_cells[i].text = str(column)
    
    # Ajouter les données du DataFrame
    for index, row in grouped_df.iterrows():
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Ajouter une image ou une autre section si nécessaire
    # doc.add_picture('path/to/image.png', width=Inches(1.0))

    # Enregistrer le document
    doc.save('rapport_donnees1.docx')

# Assurez-vous d'appeler cette fonction après le traitement des données

# Fonction pour créer le document Word
def create_word_report(df, df_export, df_import, df_third_party, final_df):
    doc = Document()
    doc.add_heading('Rapport des Données', level=1)

    # Fonction pour ajouter un tableau avec un bon design
    def add_styled_table(doc, data, title):
        doc.add_heading(title, level=2)
        table = doc.add_table(rows=1, cols=len(data.columns))
        hdr_cells = table.rows[0].cells

        # Ajouter les en-têtes de colonnes
        for i, column in enumerate(data.columns):
            hdr_cells[i].text = column
            hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(12)

            # Appliquer une couleur de fond
            shading = OxmlElement('w:shd')
            shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'CCCCCC')
            hdr_cells[i]._element.get_or_add_tcPr().append(shading)

        # Ajouter les données
        for index, row in data.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                row_cells[i].paragraphs[0].runs[0].font.size = Pt(11)

                # Appliquer une couleur de fond aux cellules
                shading = OxmlElement('w:shd')
                shading.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'FFFFFF')
                row_cells[i]._element.get_or_add_tcPr().append(shading)

        # Ajouter des bordures aux cellules
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._element.get_or_add_tcPr()
                if tc_pr is not None:
                    borders = OxmlElement('w:tcBorders')
                    for border in ['top', 'left', 'bottom', 'right']:
                        border_element = OxmlElement(f'w:{border}')
                        # Utiliser l'espace de noms correct
                        border_element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', 'single')
                        borders.append(border_element)
                    tc_pr.append(borders)

    # Ajouter le tableau des données export
    add_styled_table(doc, df_export, 'Données Export')

    # Ajouter un graphique des données export
    fig_export = px.bar(df_export, x='VILLE', y='PRIX PROPOSE', title='Données Export')
    img_export = io.BytesIO(fig_export.to_image(format="png"))
    doc.add_heading('Graphique des Données Export', level=2)
    doc.add_picture(img_export, width=Inches(5))

    # Ajouter le tableau des données import
    add_styled_table(doc, df_import, 'Données Import')

    # Ajouter un graphique des données import
    fig_import = px.bar(df_import, x='VILLE', y='PRIX PROPOSE', title='Données Import')
    img_import = io.BytesIO(fig_import.to_image(format="png"))
    doc.add_heading('Graphique des Données Import', level=2)
    doc.add_picture(img_import, width=Inches(5))

    # Ajouter un tableau final avec les totaux
    add_styled_table(doc, final_df, 'Données Finales')

    # Enregistrer le document
    doc.save('rapport_donnees.docx')
# قائمة التنقل
with st.sidebar:
    selected = option_menu(
        menu_title="التنقل",  # عنوان القائمة
        options=["الصفحة 1", "الصفحة 2", "الصفحة 3", "الصفحة 4"],  # أسماء الصفحات
        icons=["house", "info", "gear","house"],  # أيقونات (اختياري)
        menu_icon="cast",  # أيقونة القائمة
        default_index=0,  # الصفحة الافتراضية
    )

# عرض الصفحات بناءً على الاختيار
if selected == "الصفحة 1":
    st.title("مرحباً بك في الصفحة 1")
    import io
    from matplotlib import pyplot as plt
    import streamlit as st
    import plotly.express as px
    import pandas as pd
    import tempfile
    import matplotlib.pyplot as plt
    import streamlit as st
    from st_aggrid import AgGrid, GridOptionsBuilder



    # Titre de l'application
    st.title("Analyse OSTP")

    # Charge le fichier Excel
    uploaded_file = st.file_uploader("Telechargie le fichier Excel ")

    if uploaded_file:
        # Lire des donnees
        df = pd.read_excel(uploaded_file)
        # Tout en majuscules et suprimie les espase a debeu pour uniformité
        df.columns = df.columns.str.strip().str.upper() 
        df['STATUT'] = df['STATUT'].str.strip().str.upper()
        df['TYPE D\'OPÉRATION'] = df['TYPE D\'OPÉRATION'].str.strip().str.upper()
        df['VILLE'] = df['VILLE'].str.strip().str.upper()  # Convertir les noms des villes en majuscules
        df['COMMERCIAL'] = df['COMMERCIAL'].str.strip().str.upper() 
        
        # Affichie les donnees originales
        st.write("les donnees originales")
        st.dataframe(df)

    # Ajouter un bouton pour télécharger les données traitées
        st.write("### Télécharger les données traitées")
        output = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            df.to_excel(writer, index=False, sheet_name='Données_Traitées')
        output.seek(0)  # Revenir au début du fichier
        
        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données en Excel",
            data=output,
            file_name="donnees_traitees.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Filtrer les données selon 'STATUT' et 'TYPE D\'OPÉRATION'
        df_filterd = df[df['STATUT']=='OFFRE ACCEPTÉE']
        df_export = df_filterd[df_filterd['TYPE D\'OPÉRATION']=='EXPORT']
        df_import = df_filterd[df_filterd['TYPE D\'OPÉRATION']=='IMPORT']
        df_third_party = df_filterd[df_filterd['TYPE D\'OPÉRATION']=='THIRD PARTY']


        # Afficher les données Export
        st.write("### Données Export :")
        st.dataframe(df_export)

        # Ajouter un bouton pour télécharger les données Export
        st.write("### Télécharger les données Export")
        output_Export = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output_Export, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            df_export.to_excel(writer, index=False, sheet_name='Données_Export')
        output_Export.seek(0)  # Revenir au début du fichier
        
        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données en Excel",
            data=output_Export,
            file_name="donnees_Export.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
        # Afficher les données Import
        st.write("### Données Import :")
        st.dataframe(df_import)

        # Ajouter un bouton pour télécharger les données Export
        st.write("### Télécharger les données Import")
        output_Import = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output_Import, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            df_import.to_excel(writer, index=False, sheet_name='Données_Export')
        output_Import.seek(0)  # Revenir au début du fichier
        
        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données en Excel",
            data=output_Import,
            file_name="donnees_Import.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        
        # Afficher les données third party
        st.write("### Données third party :")
        st.dataframe(df_third_party)

        # Ajouter un bouton pour télécharger les données third party
        st.write("### Télécharger les données third party")
        output_third_party = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output_third_party, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            df_third_party.to_excel(writer, index=False, sheet_name='Données_Export')
        output_third_party.seek(0)  # Revenir au début du fichier
        
        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données en Excel",
            data=output_third_party,
            file_name="donnees_third_party.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Exemple de grouping par "VILLE" et "DESTINATION" et somme du "PRIX PROPOSE"
        filtered_data = df

        filtered_data['PRIX PROPOSE'] = pd.to_numeric(filtered_data['PRIX PROPOSE'], errors='coerce').fillna(0)

        # Groupement par "VILLE" et "DESTINATION" et agrégation par somme des prix
        grouped_df = filtered_data.groupby(['VILLE', 'TYPE D\'OPÉRATION', 'STATUT']).size().unstack(fill_value=0)


        # Afficher les données third party
        st.write("### Données third party :")
        st.dataframe(grouped_df)


        # Filtrer les données pour ne garder que les lignes où le statut est "Accepté"
        accepted_data = filtered_data[filtered_data['STATUT'] == 'OFFRE ACCEPTÉE']

        # Groupement par "VILLE" et "TYPE D'OPÉRATION" et somme des PRIX_PROPOSE
        grouped_prices = accepted_data.groupby(['VILLE', 'TYPE D\'OPÉRATION'])['PRIX PROPOSE'].sum()

        # Afficher les résultats
        st.write("### Somme des prix proposés pour chaque type d'opération (Statut : Accepté) :")
        st.dataframe(grouped_prices)
    

        # إعادة تعيين الفهرس لجعل الأعمدة متاحة للدمج
        grouped_df_reset = grouped_df.reset_index()

        # Fusionner les deux DataFrames sur les colonnes 'VILLE' et 'TYPE D'OPÉRATION'
        final_df = grouped_df_reset.merge(grouped_prices, on=['VILLE', 'TYPE D\'OPÉRATION'], how='left')

        # Renommer la colonne des prix proposés pour plus de clarté
        final_df.rename(columns={'PRIX PROPOSE': 'Somme PRIX PROPOSE'}, inplace=True)

        # Ajouter la somme des lignes pour chaque ligne

        # Ajouter la somme des lignes pour chaque ligne, en excluant 'Somme PRIX PROPOSE'
        final_df['Somme_Ligne'] = final_df.drop(columns=['Somme PRIX PROPOSE'], errors='ignore').sum(axis=1, numeric_only=True)

        # Ajouter une ligne pour la somme des colonnes (total des prix proposés et autres colonnes)
        total_colonne = final_df.sum(axis=0, numeric_only=True)  # Somme de toutes les colonnes
        final_df.loc['Total_Colonne'] = total_colonne

        # Afficher les données combinées dans Streamlit
        st.write("### Données combinées avec totaux et somme des prix proposés :")
        st.dataframe(final_df)

            # Ajouter un bouton pour télécharger les données third party
        st.write("### Télécharger les données third party")
        output_final_df= io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output_final_df, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            final_df.to_excel(writer, index=False, sheet_name='Données_Export')
        output_final_df.seek(0)  # Revenir au début du fichier
        
        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données en final df",
            data=output_final_df,
            file_name="donnees_final_df.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )


        # Ajouter un bouton pour télécharger toutes les données dans un seul fichier Excel
        st.write("### Télécharger toutes les données dans un seul fichier Excel")
        output_all_data = io.BytesIO()  # Création d'un objet binaire en mémoire

        with pd.ExcelWriter(output_all_data, engine='xlsxwriter') as writer:
            # Écrire chaque DataFrame dans une feuille distincte
            df.to_excel(writer, index=False, sheet_name='Données_Originales')
            df_export.to_excel(writer, index=False, sheet_name='Données_Export')
            df_import.to_excel(writer, index=False, sheet_name='Données_Import')
            df_third_party.to_excel(writer, index=False, sheet_name='Données_Third_Party')
            grouped_df.to_excel(writer, sheet_name='Groupement_Ville_Statut')
            grouped_prices.to_excel(writer, sheet_name='Somme_Prix_Acceptés')
            final_df.to_excel(writer, index=False, sheet_name='Données_Finales')

        output_all_data.seek(0)  # Revenir au début du fichier

        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger toutes les données en Excel",
            data=output_all_data,
            file_name="donnees_completes.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Assurez-vous que le nom de la colonne 'DATE' est correct
        df['DATE'] = pd.to_datetime(df['DATE'], errors='coerce')  # Convertir en format datetime
        df['Mois'] = df['DATE'].dt.to_period('M')  # Extraire le mois

        grouped_data = df.groupby(['Mois', 'TYPE D\'OPÉRATION']).size().unstack(fill_value=0)

        # Afficher les données agrégées
        st.subheader("Données agrégées par Mois et Type d'Opération")
        st.write(grouped_data)


        # Convertir grouped_data en un format long
        grouped_data_long = grouped_data.reset_index().melt(
            id_vars=['Mois'],       # Colonnes à garder telles quelles
            var_name="TYPE D'OPÉRATION",  # Nom de la colonne pour les types d'opérations
            value_name='Nombre'     # Nom de la colonne pour les valeurs
        )
        grouped_data_long['Mois'] = grouped_data_long['Mois'].dt.to_timestamp()

        # Créer un graphique linéaire avec les données transformées
        fig = px.line(
            data_frame=grouped_data_long,   # Données pour le graphique
            x='Mois',                      # Mois sur l'axe X
            y='Nombre',                    # Nombre d'opérations sur l'axe Y
            color="TYPE D'OPÉRATION",      # Différencier par type d'opération
            title="Évolution des Types d'Opérations par Mois",  # Titre du graphique
            labels={'Nombre': 'Nombre d\'Opérations', 'Mois': 'Mois'}  # Libellés des axes
        )

        # Afficher le graphique dans Streamlit
        st.plotly_chart(fig)


        df_acceptee = df[df['STATUT'] == 'OFFRE ACCEPTÉE']


        # Calculer les sommes des prix proposés par mois et par source (Export, Import, Third Party)
        prix_per_month_and_source = df_acceptee.groupby(['Mois', 'TYPE D\'OPÉRATION'])['PRIX PROPOSE'].sum().unstack(fill_value=0)

        # Affichage des résultats sous forme de table
        st.write("Table des prix proposés par mois et par source des offre acceptee")
        st.write(prix_per_month_and_source)

        prix_per_month_and_source_long = prix_per_month_and_source.reset_index().melt(
        id_vars=['Mois'],       # Colonnes à garder telles quelles
        var_name='Source',      # Nom de la colonne pour les sources
        value_name='Prix'       # Nom de la colonne pour les valeurs
    )

        # Convertir la colonne 'Mois' pour qu'elle soit JSON sérialisable
        prix_per_month_and_source_long['Mois'] = prix_per_month_and_source_long['Mois'].astype(str)

        # Créer un graphique linéaire avec les données transformées
        fig = px.line(
            data_frame=prix_per_month_and_source_long,  # Données pour le graphique
            x='Mois',                                  # Mois sur l'axe X
            y='Prix',                                  # Prix sur l'axe Y
            color='Source',                            # Différencier par source
            title="Évolution des Prix Proposés par Mois et par Source",  # Titre du graphique
            labels={'Prix': 'Prix Proposé', 'Mois': 'Mois'}             # Libellés des axes
        )

        # Afficher le graphique dans Streamlit
        st.plotly_chart(fig)

        # افترض أن لديك DataFrame يسمى df
        grouped = df.groupby(['VILLE', 'TYPE D\'OPÉRATION', 'STATUT']).size().unstack(fill_value=0)

        # تحويل البيانات إلى شكل مناسب لـ Plotly (إعادة ترتيب الأعمدة)
        grouped_reset = grouped.reset_index().melt(
            id_vars=['VILLE', 'TYPE D\'OPÉRATION'], 
            var_name='STATUT', 
            value_name='Nombre d\'Opérations'
        )

        # عرض الجدول الكامل
        st.subheader("Tableau des Types d'Opérations par Ville")
        st.dataframe(grouped_reset)  # عرض الجدول مع إمكانية التفاعل (فرز، بحث...)

        # إضافة عنصر تحديد نوع العملية
        operation_type = st.selectbox(
            "Sélectionnez le type d'opération à afficher :",
            options=grouped_reset['TYPE D\'OPÉRATION'].unique(),
            index=0  # الخيار الافتراضي
        )

        # تصفية البيانات بناءً على الاختيار
        filtered_data = grouped_reset[grouped_reset['TYPE D\'OPÉRATION'] == operation_type]

        # إنشاء مخطط شريطي تفاعلي بناءً على البيانات المصفاة
        st.subheader(f"Diagramme des {operation_type} par Ville (Interactivité)")
        fig = px.bar(
            filtered_data, 
            x='VILLE', 
            y='Nombre d\'Opérations', 
            color='STATUT', 
            barmode='group',  # يمكن تغييره إلى 'stack' إذا أردت أعمدة مكدسة
            title=f"Nombre de {operation_type} par Ville (Interactivité)",
            labels={
                "VILLE": "Ville",
                "Nombre d'Opérations": "Nombre d'Opérations",
                "STATUT": "Statut"
            },
            color_discrete_sequence=px.colors.qualitative.Set1  # تخصيص الألوان
        )

        # تخصيص التفاعلية
        fig.update_layout(
            xaxis=dict(title='Ville', tickangle=-45),  # زاوية النصوص في المحور X
            yaxis=dict(title="Nombre d'Opérations"),  # تسمية المحور Y
            legend_title="Statut",  # عنوان الأسطورة
            height=600,  # ارتفاع الرسم
            bargap=0.2,  # فراغات بين الأعمدة
        )

        # عرض الرسم باستخدام Streamlit
        st.plotly_chart(fig, use_container_width=True)


        df_ac=df[df['STATUT'] == 'OFFRE ACCEPTÉE']
        df_ac['PRIX PROPOSE'] = pd.to_numeric(df_ac['PRIX PROPOSE'], errors='coerce').fillna(0)

        accepted = df_ac.groupby(['VILLE', 'COMMERCIAL'])['PRIX PROPOSE'].sum().reset_index()
        st.subheader("Données agrégées par Mois et Type d'Opération")
        st.write(accepted)
            # Créer un graphique en barres avec Plotly Express
        fig = px.bar(
            accepted,
            x='VILLE',                  # Ville sur l'axe X
            y='PRIX PROPOSE',           # Somme des prix proposés sur l'axe Y
            color='COMMERCIAL',         # Différencier les commerciaux par couleur
            barmode='group',            # Affichage des barres groupées
            labels={'PRIX PROPOSE': 'Somme des Prix Proposés', 'VILLE': 'Ville', 'COMMERCIAL': 'Commercial'},  # Libellés
            title="Somme des Prix Proposés par Ville et Commercial"  # Titre du graphique
        )

        # Afficher le graphique dans Streamlit
        st.plotly_chart(fig)

        
        grouped_data_STATU = df.groupby(['Mois', 'STATUT']).size().unstack(fill_value=0)

        # Afficher les données agrégées
        st.subheader("Données agrégées par Mois et STATUT")
        st.write(grouped_data_STATU)


        # Convertir grouped_data en un format long
        grouped_data_long = grouped_data_STATU.reset_index().melt(
            id_vars=['Mois'],       # Colonnes à garder telles quelles
            var_name="STATUT",  # Nom de la colonne pour les types d'opérations
            value_name='Nombre'     # Nom de la colonne pour les valeurs
        )
        grouped_data_long['Mois'] = grouped_data_long['Mois'].dt.to_timestamp()

        # Créer un graphique linéaire avec les données transformées
        fig = px.line(
            data_frame=grouped_data_long,   # Données pour le graphique
            x='Mois',                      # Mois sur l'axe X
            y='Nombre',                    # Nombre d'opérations sur l'axe Y
            color="STATUT",      # Différencier par type d'opération
            title="Évolution des STATUT par Mois",  # Titre du graphique
            labels={'Nombre': 'STATUT', 'Mois': 'Mois'}  # Libellés des axes
        )

        # Afficher le graphique dans Streamlit
        st.plotly_chart(fig)

        fd_com=df.groupby(['COMMERCIAL','STATUT']).size().unstack(fill_value=0)
        # Afficher les données agrégées
        st.subheader("Données agrégées par COMMERCIAL et STATUT")
        st.write(fd_com)


            # تحويل البيانات إلى شكل مناسب لـ Plotly
        fd_com_reset = fd_com.reset_index().melt(
            id_vars=['COMMERCIAL'], 
            var_name='STATUT', 
            value_name='Nombre d\'Opérations'
        )

        # إنشاء المخطط الشريطي التفاعلي
        st.subheader("Diagramme des COMMERCIAL par STATUT")
        fig = px.bar(
            fd_com_reset, 
            x='COMMERCIAL', 
            y='Nombre d\'Opérations', 
            color='STATUT', 
            barmode='group',  # يمكن تغييره إلى 'stack' للحصول على أعمدة مكدسة
            title="Nombre d'Opérations par COMMERCIAL et STATUT",
            labels={
                "COMMERCIAL": "Commercial",
                "Nombre d'Opérations": "Nombre d'Opérations",
                "STATUT": "Statut"
            },
            color_discrete_sequence=px.colors.qualitative.Set2  # تخصيص الألوان
        )

        # تخصيص مظهر المخطط
        fig.update_layout(
            xaxis=dict(title='Commercial', tickangle=-45),  # تدوير النصوص في المحور X
            yaxis=dict(title="Nombre d'Opérations"),  # تسمية المحور Y
            legend_title="Statut",  # عنوان الأسطورة
            height=600,  # ارتفاع الرسم
            bargap=0.2  # فراغ بين الأعمدة
        )
            # Afficher le graphique dans Streamlit
        st.plotly_chart(fig)



        # Titre de l'application
        st.title("Tableau de Données avec Filtrage")

        # Filtrage par COMMERCIAL
        selected_ville = st.selectbox("Filtrer par COMMERCIAL", df["COMMERCIAL"].unique())
        df_filtered = df[df["COMMERCIAL"] == selected_ville]

        # Filtrage par STATUT
        selected_operation = st.selectbox("Filtrer par STATUT", df["STATUT"].unique())
        df_filtered = df_filtered[df_filtered["STATUT"] == selected_operation]

        # Filtrage par STATUT
        selected_operation = st.selectbox("Filtrer par Mois", df["Mois"].unique())
        df_filtered = df_filtered[df_filtered["Mois"] == selected_operation]

        # Affichage du tableau filtré
        st.dataframe(df_filtered)
    # Créer le rapport Word après le traitement
    create_word_report(df, df_export, df_import, df_third_party, final_df)

        # Bouton pour télécharger le rapport
    with open('rapport_donnees.docx', 'rb') as f:
        st.download_button('Télécharger le Rapport', f, 'rapport_donnees.docx')
    
elif selected == "الصفحة 2":
    st.title("مرحباً بك في الصفحة 2")
    import io
    import pandas as pd
    import streamlit as st
    from st_aggrid import AgGrid, GridOptionsBuilder

    # Titre de l'application
    st.title("Nombre de OC par ville/ecomercial et secteur et CA par ville")

    # Charge le fichier Excel
    uploaded_file = st.file_uploader("Téléchargez le fichier Excel")

    if uploaded_file:
        # Charger toutes les feuilles du fichier Excel
        sheets = pd.ExcelFile(uploaded_file).sheet_names
        st.write("Feuilles disponibles :", sheets)
        
        # Permettre à l'utilisateur de sélectionner une feuille
        selected_sheet = st.selectbox("Sélectionnez une feuille :", sheets)
        
        # Lire les données de la feuille sélectionnée
        df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
        
        # Nettoyage des colonnes pour uniformité
        df.columns = df.columns.str.strip().str.upper()
        
        # Sélectionner uniquement les colonnes nécessaires
        columns_needed = ['SAP', 'ICPC', 'CLIENTS', 'MOTIFE', 'COMMERCIAL', 'CA', 'VILLE']
        dfC = df[columns_needed]
        
        # Vérifier et nettoyer certaines colonnes si elles existent
        if 'COMMERCIAL' in dfC.columns:
            dfC['COMMERCIAL'] = dfC['COMMERCIAL'].str.strip().str.upper()
        if 'MOTIFE' in df.columns:
            dfC['MOTIFE'] = dfC['MOTIFE'].str.strip().str.upper()
        if 'VILLE' in dfC.columns:
            dfC['VILLE'] = dfC['VILLE'].str.strip().str.upper()
        if 'CA' in dfC.columns:
            dfC['CA'] = dfC['CA'].apply(pd.to_numeric, errors='coerce')  # Convertir en numérique si nécessaire
        
        # Afficher les données originales
        st.write("Les données originales")
        st.dataframe(dfC)
        
        # Filtrer par ville
        unique_villes = dfC['VILLE'].unique()
        selected_ville = st.selectbox("Filtrer par ville", unique_villes)
        
        # Filtrer le DataFrame par la ville sélectionnée
        filtered_df = dfC[dfC['VILLE'] == selected_ville]
        
        # Calculer la somme du CA pour la ville sélectionnée
        total_ca_ville = filtered_df['CA'].sum()
        
        # Afficher les données filtrées pour la ville sélectionnée
        st.write(f"### Données filtrées pour {selected_ville}")
        st.dataframe(filtered_df)
        
        # Afficher le total du CA pour la ville sélectionnée
        st.write(f"### Total du CA pour {selected_ville} : {total_ca_ville}")
        
        # Ajouter le total du CA à la fin du DataFrame filtré
        total_row = pd.DataFrame([{
            'SAP': 'Total',
            'ICPC': '',
            'CLIENTS': '',
            'MOTIFE': '',
            'COMMERCIAL': '',
            'CA': total_ca_ville,
            'VILLE': selected_ville
        }])
        filtered_df = pd.concat([filtered_df, total_row], ignore_index=True)
        
        # Télécharger un fichier Excel avec les données filtrées par ville
        st.write(f"### Télécharger les données pour {selected_ville}")
        output_ville = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output_ville, engine='xlsxwriter') as writer:
            filtered_df.to_excel(writer, index=False, sheet_name=f'Données_{selected_ville}')
        output_ville.seek(0)  # Revenir au début du fichier
        
        # Générer un bouton pour le téléchargement
        st.download_button(
            label=f"Télécharger les données pour {selected_ville}",
            data=output_ville,
            file_name=f"donnees_{selected_ville}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
        # Supposons que df est déjà défini comme votre DataFrame d'origine
        # Groupement des données par ville
        grouped_df = df.groupby(['VILLE', 'COMMERCIAL', 'MOTIFE', 'SECTEUR']).size().unstack(fill_value=0)
        
        # Réinitialiser l'index pour inclure "VILLE", "COMMERCIAL" et "MOTIFE"
        grouped_df = grouped_df.reset_index()
        
        # Ajouter une colonne pour les totaux des lignes
        grouped_df['Total'] = grouped_df.iloc[:, 3:].sum(axis=1)
        
        # Calculer la somme des motifs 'OC' par ville
        sum_oc = df[df['MOTIFE'] == 'OC'].groupby('VILLE').size()
        
        # Ajouter une colonne pour la somme des motifs 'OC', en utilisant .get() pour éviter les erreurs de réindexation
        grouped_df['Sum_OC'] = grouped_df['VILLE'].map(sum_oc).fillna(0).astype(int)
        
        # Ajouter une ligne pour les totaux des colonnes
        totals_row = grouped_df.iloc[:, 3:].sum(axis=0)
        totals_row['VILLE'] = 'Total'
        totals_row['COMMERCIAL'] = ''
        totals_row['MOTIFE'] = ''
        totals_row['Sum_OC'] = grouped_df['Sum_OC'].sum()
        grouped_df = pd.concat([grouped_df, pd.DataFrame([totals_row])], ignore_index=True)
        
        # Assurer la cohérence des types dans toutes les colonnes
        grouped_df = grouped_df.fillna(0)
        grouped_df.iloc[:, 3:] = grouped_df.iloc[:, 3:].astype(int)
        
        # Afficher les données groupées avec pourcentages et totaux
        st.write("### Données groupées avec pourcentages et totaux :")
        st.dataframe(grouped_df)
        
        # Télécharger les données avec les colonnes "VILLE", "COMMERCIAL", et "MOTIFE"
        st.write("### Télécharger les données groupées avec 'VILLE', 'COMMERCIAL', et 'MOTIFE'")
        output_final_df = io.BytesIO()
        with pd.ExcelWriter(output_final_df, engine='xlsxwriter') as writer:
            grouped_df.to_excel(writer, index=False, sheet_name='Données_Export')
        output_final_df.seek(0)
        
        # Générer un bouton pour le téléchargement des données groupées
        st.download_button(
            label="Télécharger les données",
            data=output_final_df,
            file_name="donnees_groupées.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

        # Télécharger toutes les données des villes dans un seul fichier Excel
        st.write("### Télécharger toutes les données des villes dans un seul fichier")
        output_all_villes = io.BytesIO()
        with pd.ExcelWriter(output_all_villes, engine='xlsxwriter') as writer:
            # Pour chaque ville, ajouter une feuille avec les données filtrées
            for ville in unique_villes:
                ville_df = dfC[dfC['VILLE'] == ville]
                ville_df.to_excel(writer, index=False, sheet_name=ville)
        output_all_villes.seek(0)
        
        # Générer un bouton pour le téléchargement du fichier Excel avec toutes les villes
        st.download_button(
            label="Télécharger toutes les données",
            data=output_all_villes,
            file_name="donnees_toutes_villes.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
                        # Créer le rapport Word après le traitement
        create_word_report2(grouped_df,dfC, unique_villes)

            # Bouton pour télécharger le rapport
        with open('rapport_donnees1.docx', 'rb') as f:
            st.download_button('Télécharger le Rapport', f, 'rapport_donnees.docx')
        

elif selected == "الصفحة 3":
    st.title("مرحباً بك في الصفحة 2")
    import io
    from matplotlib import pyplot as plt
    import streamlit as st
    import plotly.express as px
    import pandas as pd
    import tempfile
    import matplotlib.pyplot as plt
    import streamlit as st
    from st_aggrid import AgGrid, GridOptionsBuilder

    # Titre de l'application
    st.title("Client Actif et Inactif")

    # Charge le fichier Excel
    uploaded_file = st.file_uploader("Téléchargez le fichier Excel")

    if uploaded_file:
        # Charger toutes les feuilles du fichier Excel
        sheets = pd.ExcelFile(uploaded_file).sheet_names
        st.write("Feuilles disponibles :", sheets)
        
        # Permettre à l'utilisateur de sélectionner une feuille
        selected_sheet = st.selectbox("Sélectionnez une feuille :", sheets)
        
        # Lire les données de la feuille sélectionnée
        df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
        
        # Nettoyage des colonnes pour uniformité
        df.columns = df.columns.str.strip().str.upper()
        
        # Vérifier et nettoyer certaines colonnes si elles existent
        if 'COMMERCIAL (E)' in df.columns:
            df['COMMERCIAL (E)'] = df['COMMERCIAL (E)'].str.strip().str.upper()
        if 'ETAT' in df.columns:
            df['ETAT'] = df['ETAT'].str.strip().str.upper()
        if 'VILLE' in df.columns:
            df['VILLE'] = df['VILLE'].str.strip().str.upper()
        
        # Afficher les données originales
        st.write("Les données originales")
        st.dataframe(df)
    # Ajouter un bouton pour télécharger les données traitées
        st.write("### Télécharger les données traitées")
        output = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            df.to_excel(writer, index=False, sheet_name='Données_Traitées')
        output.seek(0)  # Revenir au début du fichier
    
        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données en Excel",
            data=output,
            file_name="donnees_traitees.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
    # Groupement par "VILLE", "COMMERCIAL (E)" et "ETAT"
        grouped_df = df.groupby(['VILLE', 'COMMERCIAL (E)', 'ETAT']).size().unstack(fill_value=0)

        # Ajout de colonnes supplémentaires
        grouped_df['TOTAL'] = grouped_df.sum(axis=1)  # Somme des états "Actif" et "Inactif"

        if 'ACTIF' in grouped_df.columns:
            grouped_df['% ACTIF'] = (grouped_df['ACTIF'] / grouped_df['TOTAL'] * 100).round(2)  # Pourcentage des actifs
        else:
            grouped_df['% ACTIF'] = 0  # Si aucune colonne "ACTIF"

        if 'INACTIF' in grouped_df.columns:
            grouped_df['% INACTIF'] = (grouped_df['INACTIF'] / grouped_df['TOTAL'] * 100).round(2)  # Pourcentage des inactifs
        else:
            grouped_df['% INACTIF'] = 0  # Si aucune colonne "INACTIF"

        # Réorganisation des colonnes
        ordered_columns = []
        for col in grouped_df.columns:
            if col in ['ACTIF', '% ACTIF', 'INACTIF', '% INACTIF']:
                ordered_columns.append(col)
        ordered_columns.append('TOTAL')  # Ajouter 'TOTAL' en dernier

        grouped_df = grouped_df[ordered_columns]  # Réorganiser les colonnes

        # Réinitialiser l'index pour inclure "VILLE" et "COMMERCIAL (E)"
        grouped_df = grouped_df.reset_index()

        # Afficher les données groupées avec les colonnes réorganisées
        st.write("### Données groupées avec pourcentages et totaux :")
        st.dataframe(grouped_df)

        # Ajouter un bouton pour télécharger les données avec les colonnes "VILLE" et "COMMERCIAL (E)"
        st.write("### Télécharger les données groupées avec 'VILLE' et 'COMMERCIAL (E)'")
        output_final_df = io.BytesIO()  # Création d'un objet binaire en mémoire
        with pd.ExcelWriter(output_final_df, engine='xlsxwriter') as writer:  # Utiliser xlsxwriter comme moteur
            grouped_df.to_excel(writer, index=False, sheet_name='Données_Export')
        output_final_df.seek(0)  # Revenir au début du fichier

        # Générer un bouton pour le téléchargement
        st.download_button(
            label="Télécharger les données",
            data=output_final_df,
            file_name="donnees_groupées.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

                # Créer le rapport Word après le traitement
        create_word_report1(grouped_df)

            # Bouton pour télécharger le rapport
        with open('rapport_donnees1.docx', 'rb') as f:
            st.download_button('Télécharger le Rapport', f, 'rapport_donnees.docx')
        
if selected == "الصفحة 4":
    st.title("مرحباً بك في الصفحة 4")
    import io
    import pandas as pd
    import streamlit as st

    # Titre de l'application
    st.title("Analyse Pipeline par Commercial")

    # Charger le fichier Excel
    uploaded_file = st.file_uploader("Téléchargez le fichier Excel")

    if uploaded_file:
        # Charger toutes les feuilles du fichier Excel
        sheets = pd.ExcelFile(uploaded_file).sheet_names
        st.write("Feuilles disponibles :", sheets)
        
        # Sélectionner une feuille
        selected_sheet = st.selectbox("Sélectionnez une feuille :", sheets)
        
        # Lire les données de la feuille sélectionnée
        df = pd.read_excel(uploaded_file, sheet_name=selected_sheet)
        
        # Nettoyer et afficher les noms des colonnes
        df.columns = df.columns.str.strip().str.upper()
        st.write("Noms des colonnes dans le fichier :", df.columns.tolist())
        
        # Vérifier si les colonnes nécessaires existent
        required_columns = ['SALES TEAM', 'OPPORTUNITY TYPE', 'PREVIOUS PIPELINE STAGE']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            st.error(f"Les colonnes suivantes sont manquantes : {missing_columns}")
        else:
            # Nettoyer les colonnes spécifiques
            df['SALES TEAM'] = df['SALES TEAM'].str.strip().str.upper()
            df['OPPORTUNITY TYPE'] = df['OPPORTUNITY TYPE'].str.strip().str.upper()
            df['PREVIOUS PIPELINE STAGE'] = df['PREVIOUS PIPELINE STAGE'].str.strip().str.upper()
            
            # Afficher les données pour une équipe sélectionnée
            sales_teams = df['SALES TEAM'].unique()
            selected_team = st.selectbox("Sélectionnez une équipe commerciale (SALES TEAM) :", sales_teams)
            filtered_df = df[df['SALES TEAM'] == selected_team]
            
            st.write(f"### Données filtrées pour l'équipe commerciale : {selected_team}")
            st.dataframe(filtered_df)

            # Regrouper les données
            grouped_df = (
                filtered_df
                .groupby(['OPPORTUNITY TYPE', 'PREVIOUS PIPELINE STAGE'])
                .size()
                .reset_index(name='COUNT')  # Ajouter une colonne "COUNT"
            )
            
            # Ajouter une colonne 'SUM' pour la somme par OPPORTUNITY TYPE
            grouped_df['SUM'] = grouped_df.groupby('OPPORTUNITY TYPE')['COUNT'].transform('sum')
            
            # Ajouter une ligne pour le total général
            total_general = pd.DataFrame({
                'OPPORTUNITY TYPE': ['TOTAL'],
                'PREVIOUS PIPELINE STAGE': [''],
                'COUNT': [grouped_df['COUNT'].sum()],
                'SUM': [grouped_df['SUM'].sum()]
            })
            grouped_df = pd.concat([grouped_df, total_general], ignore_index=True)

            # Télécharger un fichier Excel pour l'équipe sélectionnée
            output_file = io.BytesIO()
            with pd.ExcelWriter(output_file, engine='xlsxwriter') as writer:
                # Renommer les colonnes avec le nom de la SALES TEAM
                grouped_df.columns = [f'SALES TEAM: {selected_team}' if col == 'OPPORTUNITY TYPE' else col for col in grouped_df.columns]
                grouped_df.to_excel(writer, index=False, sheet_name=selected_team)
            output_file.seek(0)
            st.download_button(
                label=f"Télécharger le fichier pour l'équipe : {selected_team}",
                data=output_file,
                file_name=f"pipeline_{selected_team}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )
            
            # Générer un fichier Excel pour toutes les équipes `SALES TEAM`
            all_teams_file = io.BytesIO()
            with pd.ExcelWriter(all_teams_file, engine='xlsxwriter') as writer:
                for team in sales_teams:
                    team_df = df[df['SALES TEAM'] == team]
                    grouped_team_df = (
                        team_df
                        .groupby(['OPPORTUNITY TYPE', 'PREVIOUS PIPELINE STAGE'])
                        .size()
                        .reset_index(name='COUNT')
         
         
                    )
                    grouped_team_df['SUM'] = grouped_team_df.groupby('OPPORTUNITY TYPE')['COUNT'].transform('sum')
                    total_row = pd.DataFrame({
                        'OPPORTUNITY TYPE': ['TOTAL'],
                        'PREVIOUS PIPELINE STAGE': [''],
                        'COUNT': [grouped_team_df['COUNT'].sum()],
                        'SUM': [grouped_team_df['SUM'].sum()]
                    })
                    grouped_team_df = pd.concat([grouped_team_df, total_row], ignore_index=True)
                    
                    # Renommer les colonnes avec le nom de la SALES TEAM
                    grouped_team_df.columns = [f'SALES TEAM: {team}' if col == 'OPPORTUNITY TYPE' else col for col in grouped_team_df.columns]
                    grouped_team_df.to_excel(writer, index=False, sheet_name=team[:31])  # Limiter à 31 caractères pour le nom de la feuille
            all_teams_file.seek(0)
            st.download_button(
                label="Télécharger le fichier Excel pour toutes les équipes",
                data=all_teams_file,
                file_name="pipeline_all_teams.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            )

                    # Créer le rapport Word après le traitement
       # Générer le rapport Word pour toutes les équipes
    create_sales_team_report(df, sales_teams)

    # Bouton pour télécharger le rapport
    with open('rapport_pipeline_teams.docx', 'rb') as f:
        st.download_button('Télécharger le Rapport des Équipes Commerciales', f, 'rapport_pipeline_teams.docx')